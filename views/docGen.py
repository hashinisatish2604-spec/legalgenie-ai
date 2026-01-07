from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =====================================================
# HELPERS
# =====================================================

def H(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def SH(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)


def P(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(10)


def S(d, k):
    return d.get(k, "________")


# =====================================================
# MAIN GENERATOR
# =====================================================

def generate_document(doc_type, d):
    doc = Document()

    # ==================================================
    # 1️⃣ BAIL APPLICATION
    # ==================================================
    if doc_type == "Bail Application":

        H(doc, "IN THE COURT OF THE HON’BLE JUDICIAL MAGISTRATE")
        P(doc, f"Jurisdiction: {S(d,'jurisdiction')}")
        P(doc, "UNDER SECTIONS 437 / 439 OF THE CODE OF CRIMINAL PROCEDURE, 1973")

        P(
            doc,
            f"The Applicant {S(d,'applicant_name')}, "
            f"S/o / D/o / W/o {S(d,'father_name')}, aged about {S(d,'age')} years, "
            f"residing at {S(d,'address')}, respectfully submits as follows:"
        )

        SH(doc, "FACTS OF THE CASE")

        P(
            doc,
            f"1. The Applicant was arrested in FIR No. {S(d,'fir_number')} "
            f"registered at {S(d,'police_station')} for offences under "
            f"{S(d,'ipc_sections')} on {S(d,'date_of_arrest')}."
        )

        P(doc, "2. The allegations are false and politically motivated.")

        SH(doc, "GROUNDS FOR BAIL")

        P(doc, "a) The applicant is innocent.")
        P(doc, "b) No previous criminal record.")
        P(doc, "c) Continued detention is unnecessary.")

        SH(doc, "PRAYER")

        P(doc, "The Applicant prays that bail may kindly be granted.")

        SH(doc, "EXECUTION")
        P(doc, "Executed in good faith.")
        P(doc, "Date: " + S(d, "date"))

    # ==================================================
    # 2️⃣ ANTICIPATORY BAIL
    # ==================================================
    elif doc_type == "Anticipatory Bail":

        H(doc, "ANTICIPATORY BAIL APPLICATION")
        P(doc, f"Court: {S(d,'court_name')}")

        P(
            doc,
            f"The Applicant {S(d,'applicant_name')} aged {S(d,'age')} years "
            f"residing at {S(d,'address')} fears arrest in FIR No. "
            f"{S(d,'fir_number')}."
        )

        SH(doc, "REASONS FOR APPREHENSION")

        P(doc, S(d, "apprehension_reason"))
        P(doc, "Previous cases: " + S(d, "previous_cases"))

        SH(doc, "PRAYER")

        P(doc, "The Applicant prays for protection under Section 438 CrPC.")

        P(doc, "Date: " + S(d, "date"))

    # ==================================================
    # 3️⃣ FIR DRAFT
    # ==================================================
    elif doc_type == "FIR Draft":

        H(doc, "FIRST INFORMATION REPORT")

        P(doc, "Police Station: " + S(d, "police_station"))
        P(doc, "District: " + S(d, "district"))

        SH(doc, "COMPLAINANT DETAILS")

        P(
            doc,
            f"Name: {S(d,'complainant_name')}, "
            f"S/o / D/o {S(d,'father_name')}, "
            f"Address: {S(d,'address')}"
        )

        SH(doc, "INCIDENT DETAILS")

        P(
            doc,
            f"Date: {S(d,'incident_date')} | Time: {S(d,'incident_time')}"
        )
        P(doc, "Place of Incident: " + S(d, "incident_place"))
        P(doc, "Facts: " + S(d, "facts"))

        P(doc, "Date: " + S(d, "date"))

    # ==================================================
    # 4️⃣ AFFIDAVIT
    # ==================================================
    elif doc_type == "Affidavit":

        H(doc, "AFFIDAVIT")

        P(
            doc,
            f"I, {S(d,'deponent_name')}, aged {S(d,'age')} years, "
            f"residing at {S(d,'address')} do hereby solemnly affirm:"
        )

        SH(doc, "STATEMENT")

        P(doc, S(d, "statement"))
        P(doc, "Purpose: " + S(d, "purpose"))

        SH(doc, "VERIFICATION")

        P(doc, S(d, "verification"))

        P(doc, "Place: " + S(d, "place"))
        P(doc, "Date: " + S(d, "date"))

    # ==================================================
    # 5️⃣ RENT AGREEMENT
    # ==================================================
    elif doc_type == "Rent Agreement":

        H(doc, "RENT AGREEMENT")

        P(
            doc,
            f"This agreement is between {S(d,'owner_name')} (Owner) "
            f"and {S(d,'tenant_name')} (Tenant)."
        )

        SH(doc, "PROPERTY DETAILS")

        P(doc, S(d, "property_address"))

        SH(doc, "FINANCIAL TERMS")

        P(doc, "Monthly Rent: ₹" + S(d, "rent_amount"))
        P(doc, "Security Deposit: ₹" + S(d, "security_deposit"))

        SH(doc, "TENURE")

        P(
            doc,
            f"From {S(d,'start_date')} to {S(d,'end_date')} "
            f"with {S(d,'notice_period')} months notice."
        )

        SH(doc, "EXECUTION")

        P(doc, "Witness 1: " + S(d, "witness_1"))
        P(doc, "Witness 2: " + S(d, "witness_2"))

    # ==================================================
    # 6️⃣ WILL
    # ==================================================
    elif doc_type == "Will":

        H(doc, "LAST WILL AND TESTAMENT")

        P(
            doc,
            f"I, {S(d,'testator_name')}, aged {S(d,'age')} years, "
            f"resident of {S(d,'address')} declare this to be my Will."
        )

        SH(doc, "ASSETS")

        P(doc, S(d, "assets"))

        SH(doc, "BENEFICIARIES")

        P(doc, S(d, "beneficiaries"))

        SH(doc, "EXECUTOR")

        P(doc, "Executor: " + S(d, "executor"))

        P(doc, "Place: " + S(d, "place"))
        P(doc, "Date: " + S(d, "date"))

    # ==================================================
    # 7️⃣ POWER OF ATTORNEY
    # ==================================================
    elif doc_type == "Power of Attorney":

        H(doc, "POWER OF ATTORNEY")

        P(
            doc,
            f"I, {S(d,'principal_name')} residing at {S(d,'principal_address')} "
            f"appoint {S(d,'agent_name')} as my lawful attorney."
        )

        SH(doc, "POWERS GRANTED")

        P(doc, S(d, "powers"))

        SH(doc, "DURATION")

        P(doc, S(d, "duration"))

        P(doc, "Place: " + S(d, "place"))
        P(doc, "Date: " + S(d, "date"))

    # ==================================================
    # 8️⃣ CUSTOM DOCUMENT
    # ==================================================
    elif doc_type == "Custom":

        H(doc, S(d, "custom_title"))
        P(doc, S(d, "custom_text"))
        P(doc, "Purpose: " + S(d, "purpose"))
        P(doc, "Jurisdiction: " + S(d, "jurisdiction"))
        P(doc, "Place: " + S(d, "place"))
        P(doc, "Date: " + S(d, "date"))

    return doc
